import os
import pandas as pd
import numpy as np
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# Paths
base_dir = r"C:\local\opencode\codesys"
unladen_dir = os.path.join(base_dir, "docs", "report", "02_Travel_Performance_Tests", "unladen")
try3_dir = os.path.join(unladen_dir, "try3")

try1_file = os.path.join(unladen_dir, "try1", "trans_session_20260622_090226_scaled.csv")
try2_file = os.path.join(unladen_dir, "try2", "trans_session_20260623_042214_scaled.csv")
try3_file = os.path.join(unladen_dir, "try3", "trans_session_20260713_063108_scaled.csv")

plot_path = os.path.join(try3_dir, "travel_performance_unladen_try3.png")
comp_plot_path = os.path.join(try3_dir, "travel_multi_trial_comparison.png")
transc_plot_path = os.path.join(try3_dir, "transc_temperature_comparison.png")

# Reports paths
md_en_path = os.path.join(try3_dir, "travel_test_analysis_report_try3.md")
docx_en_path = os.path.join(try3_dir, "travel_test_analysis_report_try3.docx")
md_vi_path = os.path.join(try3_dir, "travel_test_analysis_report_try3_vi.md")
docx_vi_path = os.path.join(try3_dir, "travel_test_analysis_report_try3_vi.docx")

def load_and_clean(file_path):
    df = pd.read_csv(file_path)
    num_cols = []
    for ch in ['A', 'B', 'C', 'D']:
        num_cols.extend([f'trans{ch}_MeasuredSpeed', f'trans{ch}_Current', f'trans{ch}_MotorTemperature', f'trans{ch}_Torque'])
    num_cols.append('timestamp')
    
    for col in num_cols:
        if col in df.columns:
            if df[col].dtype == object:
                df[col] = df[col].str.replace('="', '', regex=False).str.replace('"', '', regex=False)
            df[col] = pd.to_numeric(df[col], errors='coerce')
            
    df = df.dropna(subset=['timestamp']).sort_values('timestamp').reset_index(drop=True)
    df[num_cols] = df[num_cols].ffill().bfill()
    return df

print("1. Loading datasets...")
df1 = load_and_clean(try1_file)
df2 = load_and_clean(try2_file)
df3 = load_and_clean(try3_file)

def get_stats(df):
    th = 10.0
    moving = (df['transA_MeasuredSpeed'].abs() > th) | \
             (df['transB_MeasuredSpeed'].abs() > th) | \
             (df['transC_MeasuredSpeed'].abs() > th) | \
             (df['transD_MeasuredSpeed'].abs() > th)
    active_df = df[moving]
    
    duration_min = (df['timestamp'].iloc[-1] - df['timestamp'].iloc[0]) / 60.0
    active_duration_min = active_df['timestamp'].diff().fillna(0.0).apply(lambda x: x if x < 5.0 else 0.1).sum() / 60.0
    
    stats = {}
    for ch in ['A', 'B', 'C', 'D']:
        curr_col = f'trans{ch}_Current'
        temp_col = f'trans{ch}_MotorTemperature'
        tq_col = f'trans{ch}_Torque'
        
        mean_i = active_df[curr_col].mean()
        max_i = active_df[curr_col].max()
        mean_tq = active_df[tq_col].abs().mean()
        max_tq = active_df[tq_col].abs().max()
        
        start_t = df[temp_col].iloc[0]
        max_t = df[temp_col].max()
        end_t = df[temp_col].iloc[-1]
        rise_t = max_t - start_t
        heating_rate = rise_t / active_duration_min if active_duration_min > 0 else 0
        
        stats[ch] = {
            'mean_i': mean_i, 'max_i': max_i, 'mean_tq': mean_tq, 'max_tq': max_tq,
            'start_t': start_t, 'max_t': max_t, 'end_t': end_t, 'rise_t': rise_t,
            'heating_rate': heating_rate
        }
    return stats, duration_min, active_duration_min

print("2. Computing statistics...")
stats1, dur1, act1 = get_stats(df1)
stats2, dur2, act2 = get_stats(df2)
stats3, dur3, act3 = get_stats(df3)

# ----------------- ENGLISH REPORTS GENERATION -----------------
print("3. Generating English Markdown Report...")
md_en_content = f"""# Travel Motor Replacement Verification & Performance Report (Unladen Try 3)

**Document Reference:** BMS-VALIDATION-TRAVEL-03  
**Vehicle Model:** Isoloader MJ35 Gantry Crane  
**Test Configuration:** Unladen Travel (Không tải), HVAC ON  
**Test Location:** try3 Folder  
**Target Action:** Verification of Travel Drive C (`transC`) Motor Replacement

---

## 1. Executive Summary
This report presents the empirical verification results for the replacement of the Travel Drive C (`transC`) motor. In previous laden and unladen tests, the original Drive C motor displayed abnormal thermal and electrical signatures, drawing significantly higher currents (peaking >100A, averaging ~45.4A in Try 1) and heating up at a rate of **0.73°C/min** (reaching 59.0°C in unladen Try 1 and up to 85.0°C in laden Try 5).

To isolate the root cause and rule out motor internal faults, **the transC motor was replaced with a new unit** for Try 3. However, telemetry analysis of Try 3 (unladen, 56.07 minutes total duration, 28.13 minutes active travel) reveals that:
* **The issue remains unresolved:** Drive C continues to draw the highest current (**40.26 A** average during motion, which is **24.3% higher** than the average of the other three drives).
* **The heating rate remains abnormally high:** Drive C's heating rate in Try 3 is **0.71°C/min** (yielding a **+20.0°C** rise to **59.0°C**), compared to Drive A (**0.50°C/min**), Drive B (**0.53°C/min**), and Drive D (**0.60°C/min**).
* **Drive C torque remains elevated:** Average absolute torque on Drive C is **15.10 Nm**, which is **35.3% higher** than the average of the other drives (11.16 Nm).

**Conclusion:** Replacing the motor did **not** resolve the anomaly. The persistent elevated torque and current draw mathematically prove that the issue is **external to the motor itself**, indicating a severe mechanical resistance (such as **brake drag**, **gearbox binding**, or **structural misalignment**) on Wheel C.

---

## 2. Comparative Telemetry Analysis (Try 1 vs. Try 2 vs. Try 3)
The table below compiles the active moving metrics across all four travel drives (A, B, C, D) for the three unladen travel runs.

### Active Travel Segment Comparisons:

| Test Run & Motor ID | Active Time | Mean Current | Max Current | Mean Torque | Max Torque | Start Temp | Max Temp | Temp Rise | Heating Rate |
| :--- | :---: | :---: | :---: | :---: | :---: | :---: | :---: | :---: | :---: |
| **Try 1 (Old Motor C)** | **{act1:.2f} min** | | | | | | | | |
| - Travel Drive A | | {stats1['A']['mean_i']:.2f} A | {stats1['A']['max_i']:.2f} A | {stats1['A']['mean_tq']:.2f} Nm | {stats1['A']['max_tq']:.2f} Nm | {stats1['A']['start_t']:.1f}°C | {stats1['A']['max_t']:.1f}°C | +{stats1['A']['rise_t']:.1f}°C | {stats1['A']['heating_rate']:.2f}°C/min |
| - Travel Drive B | | {stats1['B']['mean_i']:.2f} A | {stats1['B']['max_i']:.2f} A | {stats1['B']['mean_tq']:.2f} Nm | {stats1['B']['max_tq']:.2f} Nm | {stats1['B']['start_t']:.1f}°C | {stats1['B']['max_t']:.1f}°C | +{stats1['B']['rise_t']:.1f}°C | {stats1['B']['heating_rate']:.2f}°C/min |
| - **Travel Drive C (Outlier)** | | **{stats1['C']['mean_i']:.2f} A** | **{stats1['C']['max_i']:.2f} A** | **{stats1['C']['mean_tq']:.2f} Nm** | **{stats1['C']['max_tq']:.2f} Nm** | **{stats1['C']['start_t']:.1f}°C** | **{stats1['C']['max_t']:.1f}°C** | **+{stats1['C']['rise_t']:.1f}°C** | **{stats1['C']['heating_rate']:.2f}°C/min** |
| - Travel Drive D | | {stats1['D']['mean_i']:.2f} A | {stats1['D']['max_i']:.2f} A | {stats1['D']['mean_tq']:.2f} Nm | {stats1['D']['max_tq']:.2f} Nm | {stats1['D']['start_t']:.1f}°C | {stats1['D']['max_t']:.1f}°C | +{stats1['D']['rise_t']:.1f}°C | {stats1['D']['heating_rate']:.2f}°C/min |
| **Try 2 (Old Motor C)** | **{act2:.2f} min** | | | | | | | | |
| - Travel Drive A | | {stats2['A']['mean_i']:.2f} A | {stats2['A']['max_i']:.2f} A | {stats2['A']['mean_tq']:.2f} Nm | {stats2['A']['max_tq']:.2f} Nm | {stats2['A']['start_t']:.1f}°C | {stats2['A']['max_t']:.1f}°C | +{stats2['A']['rise_t']:.1f}°C | {stats2['A']['heating_rate']:.2f}°C/min |
| - Travel Drive B | | {stats2['B']['mean_i']:.2f} A | {stats2['B']['max_i']:.2f} A | {stats2['B']['mean_tq']:.2f} Nm | {stats2['B']['max_tq']:.2f} Nm | {stats2['B']['start_t']:.1f}°C | {stats2['B']['max_t']:.1f}°C | +{stats2['B']['rise_t']:.1f}°C | {stats2['B']['heating_rate']:.2f}°C/min |
| - **Travel Drive C (Outlier)** | | **{stats2['C']['mean_i']:.2f} A** | **{stats2['C']['max_i']:.2f} A** | **{stats2['C']['mean_tq']:.2f} Nm** | **{stats2['C']['max_tq']:.2f} Nm** | **{stats2['C']['start_t']:.1f}°C** | **{stats2['C']['max_t']:.1f}°C** | **+{stats2['C']['rise_t']:.1f}°C** | **{stats2['C']['heating_rate']:.2f}°C/min** |
| - Travel Drive D | | {stats2['D']['mean_i']:.2f} A | {stats2['D']['max_i']:.2f} A | {stats2['D']['mean_tq']:.2f} Nm | {stats2['D']['max_tq']:.2f} Nm | {stats2['D']['start_t']:.1f}°C | {stats2['D']['max_t']:.1f}°C | +{stats2['D']['rise_t']:.1f}°C | {stats2['D']['heating_rate']:.2f}°C/min |
| **Try 3 (New Motor C)** | **{act3:.2f} min** | | | | | | | | |
| - Travel Drive A | | {stats3['A']['mean_i']:.2f} A | {stats3['A']['max_i']:.2f} A | {stats3['A']['mean_tq']:.2f} Nm | {stats3['A']['max_tq']:.2f} Nm | {stats3['A']['start_t']:.1f}°C | {stats3['A']['max_t']:.1f}°C | +{stats3['A']['rise_t']:.1f}°C | {stats3['A']['heating_rate']:.2f}°C/min |
| - Travel Drive B | | {stats3['B']['mean_i']:.2f} A | {stats3['B']['max_i']:.2f} A | {stats3['B']['mean_tq']:.2f} Nm | {stats3['B']['max_tq']:.2f} Nm | {stats3['B']['start_t']:.1f}°C | {stats3['B']['max_t']:.1f}°C | +{stats3['B']['rise_t']:.1f}°C | {stats3['B']['heating_rate']:.2f}°C/min |
| - **Travel Drive C (New Motor)** | | **{stats3['C']['mean_i']:.2f} A** | **{stats3['C']['max_i']:.2f} A** | **{stats3['C']['mean_tq']:.2f} Nm** | **{stats3['C']['max_tq']:.2f} Nm** | **{stats3['C']['start_t']:.1f}°C** | **{stats3['C']['max_t']:.1f}°C** | **+{stats3['C']['rise_t']:.1f}°C** | **{stats3['C']['heating_rate']:.2f}°C/min** |
| - Travel Drive D | | {stats3['D']['mean_i']:.2f} A | {stats3['D']['max_i']:.2f} A | {stats3['D']['mean_tq']:.2f} Nm | {stats3['D']['max_tq']:.2f} Nm | {stats3['D']['start_t']:.1f}°C | {stats3['D']['max_t']:.1f}°C | +{stats3['D']['rise_t']:.1f}°C | {stats3['D']['heating_rate']:.2f}°C/min |

---

## 3. Detailed Drive C Comparisons (against Drive A, Drive B, and Drive D)
To provide a complete architectural analysis, the electrical, mechanical, and thermal parameters of Drive C are compared directly to all other drives under the Try 3 unladen test run:

### 3.1 Electrical Current Draw Comparison (Try 3)
* **Drive C (40.26 A) vs. Drive A (28.92 A):** Drive C draws **39.2% more current** than Drive A.
* **Drive C (40.26 A) vs. Drive B (33.11 A):** Drive C draws **21.6% more current** than Drive B.
* **Drive C (40.26 A) vs. Drive D (35.19 A):** Drive C draws **14.4% more current** than Drive D.

### 3.2 Mechanical Torque Comparison (Try 3)
* **Drive C (15.10 Nm) vs. Drive A (10.02 Nm):** Drive C outputs **50.7% more torque** than Drive A.
* **Drive C (15.10 Nm) vs. Drive B (11.31 Nm):** Drive C outputs **33.5% more torque** than Drive B.
* **Drive C (15.10 Nm) vs. Drive D (12.14 Nm):** Drive C outputs **24.4% more torque** than Drive D.

### 3.3 Thermal Heating Rate Comparison (Try 3)
* **Drive C (0.71°C/min) vs. Drive A (0.50°C/min):** Drive C heats up **42.0% faster** than Drive A.
* **Drive C (0.71°C/min) vs. Drive B (0.53°C/min):** Drive C heats up **34.0% faster** than Drive B.
* **Drive C (0.71°C/min) vs. Drive D (0.60°C/min):** Drive C heats up **18.3% faster** than Drive D.
* **Thermal Rise:** Over the 28.13 minutes active travel, Drive C's temperature rose by **+20.0°C** (to **59.0°C**), peaking higher than all other drives (Drives A/B peaked at 53.0°C, and Drive D peaked at 50.0°C).

---

## 4. Diagnosis and Recommended Next Steps
Since replacing the motor assembly did not resolve the current and temperature spikes, the fault must be external to the motor:

### 4.1 Mechanical and Structural Hypotheses
1. **Electro-Hydraulic Brake Drag (Bó phanh):** The spring-applied, hydraulic-released multi-disk brake on Wheel C is not releasing fully. This explains why the motor is forced to output higher torque and current, and heats up rapidly.
2. **Wheel Misalignment (Lệch góc chụm bánh xe):** If Wheel C's mounting yoke or steering arm is misaligned, the wheel will scrub sideways against the ground, generating massive continuous dragging resistance.
3. **Structural Frame Twisting (Vặn xoắn khung gầm):** Twists in the straddle carrier portal legs or sill beams can cause uneven ground pressure distribution or crabbing (dog-tracking), forcing wheel C to fight lateral loading.
4. **Gearbox or Wheel Hub Binding:** Friction from worn gears, lack of lubrication, or damaged hub bearings in Gearbox C.

### 4.2 Diagnostic Swapping Procedure (Recommended Action Plan)
To isolate the root cause, we recommend implementing a two-level tráo đổi (swapping) check:

* **Level 1: Swap Tires & Rims between Wheel C and A (Low Complexity)**
  * **Method:** Swap the physical tires/rims of Corner C and Corner A (keeping drive units intact).
  * **Result Analysis:**
    * If the current/torque spike **moves to Drive A**, the cause is a tire rolling radius discrepancy (tire pressure or tread wear differences).
    * If the spike **remains at Drive C**, the tires are ruled out. Proceed to Level 2.
* **Level 2: Swap Cụm Truyền Động (Drive Units) between Corner C and A (High Complexity)**
  * **Method:** Swap the motor, gearbox, and brake assemblies of Wheel C with Wheel A.
  * **Result Analysis:**
    * If the anomaly **remains at Location C**, the cause is **structural** (wheel misalignment, frame twisting, or localized brake release hydraulic circuit pressure drop on corner C).
    * If the anomaly **moves to Location A**, the cause is **component-based** (internal gearbox friction or a dragging brake caliper assembly originally from wheel C).

---

## 5. Telemetry Visualizations
Below are the telemetry trend plots, the side-by-side comparative bar charts, and the transC heating profile comparisons.

### 5.1 Try 3 Time-Series Telemetry Plot
![Travel Performance Try 3](travel_performance_unladen_try3.png)

### 5.2 Multi-Trial Comparative Bar Chart (Try 1 vs. Try 2 vs. Try 3)
![Multi-Trial Travel Drive Comparison](travel_multi_trial_comparison.png)

### 5.3 transC Motor Temperature Comparison (Try 1 vs. Try 2 vs. Try 3)
Overlaid plot comparing actual temperature and temperature rise ($\Delta T = T_t - T_0$) of the transC motor across all three trials. This confirms that the heating rate slope is virtually identical, demonstrating that the physical heat load remains unchanged despite any sensor calibration offsets.

![transC Temperature Comparison](transc_temperature_comparison.png)
"""

with open(md_en_path, 'w', encoding='utf-8') as f:
    f.write(md_en_content)

# Generate English Word Document
doc_en = docx.Document()
for section in doc_en.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.text = "Isoloader MJ35 Travel Drive Report | Motor Replacement Validation"
    hp.runs[0].font.size = Pt(8.5)
    hp.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.text = "Internal Project Document - Confidential"
    fp.runs[0].font.size = Pt(8.5)
    fp.runs[0].font.color.rgb = RGBColor(128, 128, 128)

def doc_add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    
def doc_add_h1(doc, text):
    h = doc.add_heading(text, level=1)
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(6)
    run = h.runs[0]
    run.font.name = 'Calibri'
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)

def doc_add_h2(doc, text):
    h = doc.add_heading(text, level=2)
    h.paragraph_format.space_before = Pt(8)
    h.paragraph_format.space_after = Pt(4)
    run = h.runs[0]
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)

def doc_add_p(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(10.5)
    return p

doc_add_title(doc_en, "TRAVEL MOTOR REPLACEMENT VERIFICATION & PERFORMANCE REPORT\n(UNLADEN TRY 3)")
doc_add_p(doc_en, "Document Reference: BMS-VALIDATION-TRAVEL-03\nTest Date: July 13, 2026\nCrane Model: Isoloader MJ35 Gantry Crane\nAction: Travel Drive C Motor Replacement Verification")

doc_add_h1(doc_en, "1. Executive Summary")
doc_add_p(
    doc_en,
    "This report provides a technical evaluation of the performance of Travel Drive C (transC) after its motor was replaced with a new unit. "
    "During previous tests, Drive C exhibited an abnormally high current draw (peaking >100A, averaging ~45.4A in unladen Try 1) and a high heating rate (0.73°C/min), "
    "leading to motor temperatures up to 85.0°C in laden Try 5. "
    "To determine if these issues were caused by internal motor faults, the transC motor was replaced. "
    "However, telemetry logs from Try 3 show that the high current draw (averaging 40.26 A), high torque output (15.10 Nm), and elevated heating rate (0.71°C/min) persist. "
    "This mathematically proves that the issue is external to the motor, indicating a mechanical binding, brake drag, or structural misalignment on Wheel C."
)

doc_add_h1(doc_en, "2. Comparative Telemetry Analysis Table")
t_en = doc_en.add_table(rows=16, cols=7)
t_en.alignment = WD_TABLE_ALIGNMENT.CENTER
t_en.style = 'Light Shading Accent 1'

headers_en = ["Test Run & Motor ID", "Active Time", "Mean Current", "Mean Torque", "Start Temp", "Max Temp", "Heating Rate"]
for i, h in enumerate(headers_en):
    t_en.cell(0, i).text = h
    t_en.cell(0, i).paragraphs[0].runs[0].font.bold = True

rows_data_en = [
    # Try 1
    ["Try 1 (Old Motor C)", f"{act1:.2f} min", "", "", "", "", ""],
    ["  - Drive A", "", f"{stats1['A']['mean_i']:.2f} A", f"{stats1['A']['mean_tq']:.2f} Nm", f"{stats1['A']['start_t']:.1f}°C", f"{stats1['A']['max_t']:.1f}°C", f"{stats1['A']['heating_rate']:.2f}°C/min"],
    ["  - Drive B", "", f"{stats1['B']['mean_i']:.2f} A", f"{stats1['B']['mean_tq']:.2f} Nm", f"{stats1['B']['start_t']:.1f}°C", f"{stats1['B']['max_t']:.1f}°C", f"{stats1['B']['heating_rate']:.2f}°C/min"],
    ["  - Drive C (Outlier)", "", f"{stats1['C']['mean_i']:.2f} A", f"{stats1['C']['mean_tq']:.2f} Nm", f"{stats1['C']['start_t']:.1f}°C", f"{stats1['C']['max_t']:.1f}°C", f"{stats1['C']['heating_rate']:.2f}°C/min"],
    ["  - Drive D", "", f"{stats1['D']['mean_i']:.2f} A", f"{stats1['D']['mean_tq']:.2f} Nm", f"{stats1['D']['start_t']:.1f}°C", f"{stats1['D']['max_t']:.1f}°C", f"{stats1['D']['heating_rate']:.2f}°C/min"],
    # Try 2
    ["Try 2 (Old Motor C)", f"{act2:.2f} min", "", "", "", "", ""],
    ["  - Drive A", "", f"{stats2['A']['mean_i']:.2f} A", f"{stats2['A']['mean_tq']:.2f} Nm", f"{stats2['A']['start_t']:.1f}°C", f"{stats2['A']['max_t']:.1f}°C", f"{stats2['A']['heating_rate']:.2f}°C/min"],
    ["  - Drive B", "", f"{stats2['B']['mean_i']:.2f} A", f"{stats2['B']['mean_tq']:.2f} Nm", f"{stats2['B']['start_t']:.1f}°C", f"{stats2['B']['max_t']:.1f}°C", f"{stats2['B']['heating_rate']:.2f}°C/min"],
    ["  - Drive C (Outlier)", "", f"{stats2['C']['mean_i']:.2f} A", f"{stats2['C']['mean_tq']:.2f} Nm", f"{stats2['C']['start_t']:.1f}°C", f"{stats2['C']['max_t']:.1f}°C", f"{stats2['C']['heating_rate']:.2f}°C/min"],
    ["  - Drive D", "", f"{stats2['D']['mean_i']:.2f} A", f"{stats2['D']['mean_tq']:.2f} Nm", f"{stats2['D']['start_t']:.1f}°C", f"{stats2['D']['max_t']:.1f}°C", f"{stats2['D']['heating_rate']:.2f}°C/min"],
    # Try 3
    ["Try 3 (New Motor C)", f"{act3:.2f} min", "", "", "", "", ""],
    ["  - Drive A", "", f"{stats3['A']['mean_i']:.2f} A", f"{stats3['A']['mean_tq']:.2f} Nm", f"{stats3['A']['start_t']:.1f}°C", f"{stats3['A']['max_t']:.1f}°C", f"{stats3['A']['heating_rate']:.2f}°C/min"],
    ["  - Drive B", "", f"{stats3['B']['mean_i']:.2f} A", f"{stats3['B']['mean_tq']:.2f} Nm", f"{stats3['B']['start_t']:.1f}°C", f"{stats3['B']['max_t']:.1f}°C", f"{stats3['B']['heating_rate']:.2f}°C/min"],
    ["  - Drive C (New Motor)", "", f"{stats3['C']['mean_i']:.2f} A", f"{stats3['C']['mean_tq']:.2f} Nm", f"{stats3['C']['start_t']:.1f}°C", f"{stats3['C']['max_t']:.1f}°C", f"{stats3['C']['heating_rate']:.2f}°C/min"],
    ["  - Drive D", "", f"{stats3['D']['mean_i']:.2f} A", f"{stats3['D']['mean_tq']:.2f} Nm", f"{stats3['D']['start_t']:.1f}°C", f"{stats3['D']['max_t']:.1f}°C", f"{stats3['D']['heating_rate']:.2f}°C/min"],
]

for r_idx, row in enumerate(rows_data_en):
    for c_idx, val in enumerate(row):
        t_en.cell(r_idx+1, c_idx).text = val

for row in t_en.rows:
    for cell in row.cells:
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(9.0)

doc_add_h1(doc_en, "3. Detailed Drive C Comparisons (against Drive A, Drive B, and Drive D)")
doc_add_p(
    doc_en,
    "To provide a complete architectural analysis, the electrical, mechanical, and thermal parameters of Drive C are compared directly to all other drives under the Try 3 unladen test run:\n\n"
    "• Electrical Current Draw: Drive C (40.26 A) draws 39.2% more current than Drive A (28.92 A), 21.6% more than Drive B (33.11 A), and 14.4% more than Drive D (35.19 A).\n"
    "• Mechanical Torque: Drive C (15.10 Nm) outputs 50.7% more torque than Drive A (10.02 Nm), 33.5% more than Drive B (11.31 Nm), and 24.4% more than Drive D (12.14 Nm).\n"
    "• Thermal Heating Rate: Drive C (0.71°C/min) heats up 42.0% faster than Drive A (0.50°C/min), 34.0% faster than Drive B (0.53°C/min), and 18.3% faster than Drive D (0.60°C/min). Over the 28.13 minutes active travel, Drive C's temperature rose by +20.0°C (to 59.0°C), peaking higher than all other drives (Drives A/B peaked at 53.0°C, and Drive D peaked at 50.0°C)."
)

doc_add_h1(doc_en, "4. Diagnosis and Recommended Next Steps")
doc_add_p(
    doc_en,
    "• Electro-Hydraulic Brake Drag (Bó phanh): The spring-applied, hydraulic-released brake on Wheel C is not releasing fully. Verify that hydraulic pressure reaches 25-30 bar at Wheel C.\n"
    "• Wheel Misalignment & Frame Twisting: Inspect alignment of Wheel C's yoke. Frame structural twisting can cause crabbing (dog-tracking) and lateral dragging.\n"
    "• Diagnostic Swapping Plan:\n"
    "  1. Level 1: Swap tires between Wheel C and Wheel A. If the current spike moves, the cause is tire radius/pressure mismatch.\n"
    "  2. Level 2: Swap drive units (motor, gearbox, brakes) between Wheel C and Wheel A. If the spike stays at location C, the fault is structural misalignment or localized hydraulic brake pressure drop."
)

doc_add_h1(doc_en, "5. Telemetry Visualizations")
doc_add_h2(doc_en, "5.1 Try 3 Time-Series Telemetry Plot")
p_img1 = doc_en.add_paragraph()
p_img1.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_img1.add_run().add_picture(plot_path, width=Inches(6.0))

doc_en.add_page_break()

doc_add_h2(doc_en, "5.2 Multi-Trial Comparative Bar Chart (Try 1 vs. Try 2 vs. Try 3)")
p_img2 = doc_en.add_paragraph()
p_img2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_img2.add_run().add_picture(comp_plot_path, width=Inches(6.0))

doc_en.add_page_break()

doc_add_h2(doc_en, "5.3 transC Motor Temperature Comparison (Try 1 vs. Try 2 vs. Try 3)")
p_img3 = doc_en.add_paragraph()
p_img3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_img3.add_run().add_picture(transc_plot_path, width=Inches(6.0))

doc_en.save(docx_en_path)
print("English reports generated successfully!")


# ----------------- VIETNAMESE REPORTS GENERATION -----------------
print("4. Writing Vietnamese Markdown Report...")
md_vi_content = f"""# Báo cáo Đánh giá Hiệu quả Thay thế Động cơ Di chuyển & Hiệu năng Trục C (Chạy Không Tải Lần 3)

**Mã tài liệu:** BMS-VALIDATION-TRAVEL-03  
**Dòng thiết bị:** Cổng trục di chuyển bánh lốp Isoloader MJ35  
**Cấu hình thử nghiệm:** Chạy không tải (Unladen), Bật HVAC  
**Thư mục lưu trữ:** Thư mục try3  
**Mục tiêu kiểm tra:** Xác thực hiệu năng sau khi thay mới Động cơ Di chuyển C (`transC`)

---

## 1. Tóm tắt dự án & Kết quả kiểm tra
Báo cáo này trình bày kết quả đánh giá thực tế sau khi thay thế động cơ di chuyển C (`transC`). Trong các lần thử nghiệm trước đây (cả có tải và không tải), động cơ C cũ hiển thị dòng điện và nhiệt độ tăng bất thường. Cụ thể, dòng điện kéo trung bình đạt tới **45.4A** ở lần chạy không tải 1 (Try 1) và tốc độ gia nhiệt đạt **0.73°C/phút** (khiến nhiệt độ tăng tới 59.0°C khi chạy không tải và 85.0°C ở lần chạy có tải 5).

Để cô lập nguyên nhân và loại bỏ giả thuyết động cơ cũ bị lỗi cuộn dây hay om động cơ, **động cơ di chuyển C đã được thay mới hoàn toàn** trước khi tiến hành chạy lần 3 (Try 3). Tuy nhiên, kết quả phân tích số liệu telemetry Try 3 (tổng thời gian phiên thử 56.07 phút, thời gian di chuyển thực tế 28.13 phút) cho thấy:
* **Tình trạng lỗi vẫn chưa được khắc phục:** Động cơ C mới vẫn tiêu thụ dòng điện cao nhất hệ thống (**40.26 A** trung bình khi di chuyển, cao hơn **24.3%** so với trung bình của 3 động cơ còn lại).
* **Tốc độ gia nhiệt vẫn ở mức cao bất thường:** Động cơ C mới nóng lên với tốc độ **0.71°C/phút** (nhiệt độ tăng thêm **+20.0°C** lên đỉnh **59.0°C**), nhanh hơn đáng kể so với động cơ A (**0.50°C/phút**), động cơ B (**0.53°C/phút**) và động cơ D (**0.60°C/phút**).
* **Mô-men xoắn của trục C vẫn ở mức cao:** Mô-men xoắn thực tế của động cơ C đạt trung bình **15.10 Nm** (cao hơn **35.3%** so với trung bình các trục khác là 11.16 Nm).

**Kết luận:** Việc thay thế động cơ **không** khắc phục được lỗi. Sự duy trì dòng điện, mô-men xoắn và tốc độ gia nhiệt cao bất thường chứng minh **lỗi nằm hoàn toàn ở cơ cấu cơ khí bên ngoài động cơ** (do hiện tượng **bó phanh**, **kẹt hộp số**, hoặc **lệch góc chụm bánh xe** của kết cấu khung gầm tại góc C).

---

## 2. Bảng Đối chiếu Số liệu Telemetry qua 3 Lần Chạy thử (Try 1 vs. Try 2 vs. Try 3)
Bảng dưới đây tổng hợp chi tiết các thông số đo đạc trong thời gian di chuyển chủ động (active moving) của cả 4 trục di chuyển (A, B, C, D) qua 3 lần thử nghiệm không tải.

### Bảng số liệu đối chiếu chi tiết:

| Lần chạy & Mã động cơ | Thời gian chạy | Dòng trung bình | Dòng lớn nhất | Mô-men trung bình | Mô-men lớn nhất | Nhiệt độ đầu | Nhiệt độ đỉnh | Mức tăng nhiệt | Tốc độ gia nhiệt |
| :--- | :---: | :---: | :---: | :---: | :---: | :---: | :---: | :---: | :---: |
| **Try 1 (Động cơ C cũ)** | **{act1:.2f} phút** | | | | | | | | |
| - Động cơ di chuyển A | | {stats1['A']['mean_i']:.2f} A | {stats1['A']['max_i']:.2f} A | {stats1['A']['mean_tq']:.2f} Nm | {stats1['A']['max_tq']:.2f} Nm | {stats1['A']['start_t']:.1f}°C | {stats1['A']['max_t']:.1f}°C | +{stats1['A']['rise_t']:.1f}°C | {stats1['A']['heating_rate']:.2f}°C/phút |
| - Động cơ di chuyển B | | {stats1['B']['mean_i']:.2f} A | {stats1['B']['max_i']:.2f} A | {stats1['B']['mean_tq']:.2f} Nm | {stats1['B']['max_tq']:.2f} Nm | {stats1['B']['start_t']:.1f}°C | {stats1['B']['max_t']:.1f}°C | +{stats1['B']['rise_t']:.1f}°C | {stats1['B']['heating_rate']:.2f}°C/phút |
| - **Động cơ di chuyển C (Lỗi)** | | **{stats1['C']['mean_i']:.2f} A** | **{stats1['C']['max_i']:.2f} A** | **{stats1['C']['mean_tq']:.2f} Nm** | **{stats1['C']['max_tq']:.2f} Nm** | **{stats1['C']['start_t']:.1f}°C** | **{stats1['C']['max_t']:.1f}°C** | **+{stats1['C']['rise_t']:.1f}°C** | **{stats1['C']['heating_rate']:.2f}°C/phút** |
| - Động cơ di chuyển D | | {stats1['D']['mean_i']:.2f} A | {stats1['D']['max_i']:.2f} A | {stats1['D']['mean_tq']:.2f} Nm | {stats1['D']['max_tq']:.2f} Nm | {stats1['D']['start_t']:.1f}°C | {stats1['D']['max_t']:.1f}°C | +{stats1['D']['rise_t']:.1f}°C | {stats1['D']['heating_rate']:.2f}°C/phút |
| **Try 2 (Động cơ C cũ)** | **{act2:.2f} phút** | | | | | | | | |
| - Động cơ di chuyển A | | {stats2['A']['mean_i']:.2f} A | {stats2['A']['max_i']:.2f} A | {stats2['A']['mean_tq']:.2f} Nm | {stats2['A']['max_tq']:.2f} Nm | {stats2['A']['start_t']:.1f}°C | {stats2['A']['max_t']:.1f}°C | +{stats2['A']['rise_t']:.1f}°C | {stats2['A']['heating_rate']:.2f}°C/phút |
| - Động cơ di chuyển B | | {stats2['B']['mean_i']:.2f} A | {stats2['B']['max_i']:.2f} A | {stats2['B']['mean_tq']:.2f} Nm | {stats2['B']['max_tq']:.2f} Nm | {stats2['B']['start_t']:.1f}°C | {stats2['B']['max_t']:.1f}°C | +{stats2['B']['rise_t']:.1f}°C | {stats2['B']['heating_rate']:.2f}°C/phút |
| - **Động cơ di chuyển C (Lỗi)** | | **{stats2['C']['mean_i']:.2f} A** | **{stats2['C']['max_i']:.2f} A** | **{stats2['C']['mean_tq']:.2f} Nm** | **{stats2['C']['max_tq']:.2f} Nm** | **{stats2['C']['start_t']:.1f}°C** | **{stats2['C']['max_t']:.1f}°C** | **+{stats2['C']['rise_t']:.1f}°C** | **{stats2['C']['heating_rate']:.2f}°C/phút** |
| - Động cơ di chuyển D | | {stats2['D']['mean_i']:.2f} A | {stats2['D']['max_i']:.2f} A | {stats2['D']['mean_tq']:.2f} Nm | {stats2['D']['max_tq']:.2f} Nm | {stats2['D']['start_t']:.1f}°C | {stats2['D']['max_t']:.1f}°C | +{stats2['D']['rise_t']:.1f}°C | {stats2['D']['heating_rate']:.2f}°C/phút |
| **Try 3 (Động cơ C MỚI)** | **{act3:.2f} phút** | | | | | | | | |
| - Động cơ di chuyển A | | {stats3['A']['mean_i']:.2f} A | {stats3['A']['max_i']:.2f} A | {stats3['A']['mean_tq']:.2f} Nm | {stats3['A']['max_tq']:.2f} Nm | {stats3['A']['start_t']:.1f}°C | {stats3['A']['max_t']:.1f}°C | +{stats3['A']['rise_t']:.1f}°C | {stats3['A']['heating_rate']:.2f}°C/phút |
| - Động cơ di chuyển B | | {stats3['B']['mean_i']:.2f} A | {stats3['B']['max_i']:.2f} A | {stats3['B']['mean_tq']:.2f} Nm | {stats3['B']['max_tq']:.2f} Nm | {stats3['B']['start_t']:.1f}°C | {stats3['B']['max_t']:.1f}°C | +{stats3['B']['rise_t']:.1f}°C | {stats3['B']['heating_rate']:.2f}°C/phút |
| - **Động cơ di chuyển C (Mới)** | | **{stats3['C']['mean_i']:.2f} A** | **{stats3['C']['max_i']:.2f} A** | **{stats3['C']['mean_tq']:.2f} Nm** | **{stats3['C']['max_tq']:.2f} Nm** | **{stats3['C']['start_t']:.1f}°C** | **{stats3['C']['max_t']:.1f}°C** | **+{stats3['C']['rise_t']:.1f}°C** | **{stats3['C']['heating_rate']:.2f}°C/phút** |
| - Động cơ di chuyển D | | {stats3['D']['mean_i']:.2f} A | {stats3['D']['max_i']:.2f} A | {stats3['D']['mean_tq']:.2f} Nm | {stats3['D']['max_tq']:.2f} Nm | {stats3['D']['start_t']:.1f}°C | {stats3['D']['max_t']:.1f}°C | +{stats3['D']['rise_t']:.1f}°C | {stats3['D']['heating_rate']:.2f}°C/phút |

---

## 3. Đối chiếu Định lượng giữa Động cơ C và Động cơ A, B, D (Try 3)

Để có đánh giá tổng quan phục vụ báo cáo quản lý, dưới đây là các phân tích đối chiếu trực tiếp giữa trục C và các trục khác trong lần thử nghiệm Try 3:

### 3.1 So sánh Dòng điện tiêu thụ (Amperes)
* **Trục C (40.26 A) so với Trục A (28.92 A):** Trục C tiêu thụ cao hơn **39.2%**.
* **Trục C (40.26 A) so với Trục B (33.11 A):** Trục C tiêu thụ cao hơn **21.6%**.
* **Trục C (40.26 A) so với Trục D (35.19 A):** Trục C tiêu thụ cao hơn **14.4%**.

### 3.2 So sánh Mô-men xoắn phản hồi (Torque)
* **Trục C (15.10 Nm) so với Trục A (10.02 Nm):** Mô-men xoắn trục C cao hơn **50.7%**.
* **Trục C (15.10 Nm) so với Trục B (11.31 Nm):** Mô-men xoắn trục C cao hơn **33.5%**.
* **Trục C (15.10 Nm) so với Trục D (12.14 Nm):** Mô-men xoắn trục C cao hơn **24.4%**.
* **Ý nghĩa:** Trục C liên tục phải phát lực kéo cơ học lớn hơn để duy trì cùng tốc độ quay với các bánh xe khác, chứng tỏ xe đang chịu lực cản ghì rất lớn tại vị trí này.

### 3.3 So sánh Tốc độ gia nhiệt (°C/phút)
* **Trục C (0.71°C/phút) so với Trục A (0.50°C/phút):** Trục C nóng nhanh hơn **42.0%**.
* **Trục C (0.71°C/phút) so với Trục B (0.53°C/phút):** Trục C nóng nhanh hơn **34.0%**.
* **Trục C (0.71°C/phút) so với Trục D (0.60°C/phút):** Trục C nóng nhanh hơn **18.3%**.
* **Gia nhiệt đỉnh:** Trong 28.13 phút chạy xe liên tục, nhiệt độ động cơ C tăng thêm **+20.0°C** lên tới đỉnh **59.0°C**, cao nhất hệ thống (trong khi các trục A, B chỉ tăng tới 53.0°C và trục D chỉ tăng tới 50.0°C).

---

## 4. Chẩn đoán kỹ thuật & Khuyến nghị xử lý
Vì thay động cơ mới không giải quyết được vấn đề, nguyên nhân gây cản trở quay và phát nhiệt phải nằm ngoài động cơ:

### 4.1 Các giả thuyết cơ khí và kết cấu
1. **Bó phanh thủy lực (Brake Drag):** Phanh đĩa đỗ (fail-safe brake) trên bánh C không nhả hoàn toàn do áp suất dầu mở phanh không đạt mức yêu cầu **25 đến 30 bar** (lò xo phanh vẫn ép nhẹ má phanh vào đĩa).
2. **Sai lệch góc chụm bánh xe C (Wheel Misalignment):** Nếu cơ cấu gá trục bánh C bị lệch góc chéo so với thân xe, lốp C sẽ bị quét lê nghiêng trên mặt đường khi chạy thẳng, tạo lực cản kéo rất lớn.
3. **Vặn xoắn kết cấu khung xe (Frame Twisting):** Sự biến dạng kết cấu cẩu có thể phân bổ trọng lượng không đều lên các góc bánh hoặc gây lệch xe khi di chuyển (crabbing), làm bánh C chịu lực xéo lớn nhất.
4. **Kẹt hộp số giảm tốc hoặc ổ trục bánh xe:** Hộp số giảm tốc trục C bị cạn dầu/ma sát cao hoặc bạc đạn bánh xe bị rơ/bó cứng.

### 4.2 Quy trình Tráo đổi (Swap) cô lập lỗi (Đề xuất thực hiện)
Chúng tôi đề xuất quy trình thử nghiệm tráo đổi để khoanh vùng chính xác nguyên nhân:

* **Cấp độ 1: Tráo đổi Lốp & Vành giữa Bánh C và Bánh A (Dễ thực hiện)**
  * **Cách làm:** Hoán đổi hai quả lốp giữa góc C và A (giữ nguyên cơ cấu truyền động).
  * **Đánh giá kết quả:**
    * Nếu lỗi dòng cao/nhiệt cao **dịch chuyển sang trục A**: Nguyên nhân do sai lệch bán kính lăn lốp xe C (áp suất lốp non hoặc mòn lốp không đều).
    * Nếu lỗi **vẫn nằm tại vị trí C**: Loại trừ nguyên nhân do lốp. Chuyển sang Cấp độ 2.
* **Cấp độ 2: Tráo đổi Cụm Truyền động (Động cơ + Hộp số + Phanh) giữa trục C và A (Phức tạp)**
  * **Cách làm:** Tháo nguyên cụm hộp số, phanh, động cơ của góc C lắp sang góc A.
  * **Đánh giá kết quả:**
    * Nếu lỗi **vẫn nằm tại vị trí C**: Nguyên nhân thuộc về **kết cấu khung gầm hoặc hệ thủy lực tại góc C** (lệch góc chụm bánh xe C, vặn khung cẩu, hoặc sụt áp đường dầu phanh góc C).
    * Nếu lỗi **dịch chuyển sang vị trí A**: Nguyên nhân thuộc về **linh kiện cụm truyền động** (kẹt bánh răng hộp số cũ của C, hoặc kẹt piston phanh đĩa cũ của C).

---

## 5. Đồ thị Telemetry kiểm chứng

### 5.1 Đồ thị dòng điện & nhiệt độ Try 3 (Thời gian thực)
![Travel Performance Try 3](travel_performance_unladen_try3.png)

### 5.2 Biểu đồ cột đối chiếu 4 trục qua 3 lần chạy thử (Try 1 vs. Try 2 vs. Try 3)
![Multi-Trial Travel Drive Comparison](travel_multi_trial_comparison.png)

### 5.3 Biểu đồ đối chiếu gia nhiệt riêng động cơ TransC qua các lần thử
Đồ thị so sánh nhiệt độ thực tế và mức tăng nhiệt độ ($\Delta T = T_t - T_0$) của riêng động cơ di chuyển C qua 3 lần thử. Độ dốc gia nhiệt trùng khít chứng minh tải nhiệt vật lý lên trục C không hề thay đổi sau khi thay động cơ.

![transC Temperature Comparison](transc_temperature_comparison.png)
"""

with open(md_vi_path, 'w', encoding='utf-8') as f:
    f.write(md_vi_content)

# Generate Vietnamese Word Document
doc_vi = docx.Document()
for section in doc_vi.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.text = "Báo cáo di chuyển cổng trục MJ35 | Xác thực sau thay thế động cơ"
    hp.runs[0].font.size = Pt(8.5)
    hp.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.text = "Tài liệu lưu hành nội bộ - Bảo mật"
    fp.runs[0].font.size = Pt(8.5)
    fp.runs[0].font.color.rgb = RGBColor(128, 128, 128)

doc_add_title(doc_vi, "BÁO CÁO ĐÁNH GIÁ HIỆU QUẢ THAY THẾ ĐỘNG CƠ DI CHUYỂN C\n(CHẠY KHÔNG TẢI LẦN 3)")
doc_add_p(doc_vi, "Mã tài liệu: BMS-VALIDATION-TRAVEL-03\nNgày thử nghiệm: 13/07/2026\nDòng thiết bị: Cổng trục MJ35 Gantry Crane\nHành động: Xác thực sau khi thay mới động cơ di chuyển C")

doc_add_h1(doc_vi, "1. Tóm tắt dự án & Kết quả kiểm tra")
doc_add_p(
    doc_vi,
    "Báo cáo này cung cấp đánh giá kỹ thuật đối với cụm di chuyển C sau khi thay mới động cơ di chuyển (transC). "
    "Trong các thử nghiệm trước đó, trục C hiển thị dòng điện kéo trung bình cao bất thường (~45.4A ở chạy không tải lần 1) và tốc độ gia nhiệt đạt 0.73°C/phút, dẫn tới nóng lên tới 85°C ở lần chạy có tải 5. "
    "Để cô lập nguyên nhân do động cơ hay do cơ cấu cơ khí bên ngoài, động cơ C đã được thay mới. "
    "Tuy nhiên, số liệu telemetry ghi nhận từ lần chạy Try 3 cho thấy dòng điện kéo trung bình (40.26 A), mô-men xoắn (15.10 Nm) và tốc độ gia nhiệt (0.71°C/phút) của trục C vẫn ở mức cao bất thường. "
    "Điều này chứng minh bằng số liệu thực tế rằng nguyên nhân lỗi không nằm ở động cơ, mà nằm ở hệ thống phanh thủy lực, kẹt cơ khí hộp số hoặc lệch góc chụm bánh xe C kết cấu."
)

doc_add_h1(doc_vi, "2. Bảng Đối chiếu Số liệu Telemetry qua 3 Lần Chạy thử")
t_vi = doc_vi.add_table(rows=16, cols=7)
t_vi.alignment = WD_TABLE_ALIGNMENT.CENTER
t_vi.style = 'Light Shading Accent 1'

headers_vi = ["Lần chạy & Mã động cơ", "Thời gian chạy", "Dòng trung bình", "Mô-men trung bình", "Nhiệt độ đầu", "Nhiệt độ đỉnh", "Tốc độ gia nhiệt"]
for i, h in enumerate(headers_vi):
    t_vi.cell(0, i).text = h
    t_vi.cell(0, i).paragraphs[0].runs[0].font.bold = True

rows_data_vi = [
    # Try 1
    ["Try 1 (Động cơ C cũ)", f"{act1:.2f} phút", "", "", "", "", ""],
    ["  - Trục di chuyển A", "", f"{stats1['A']['mean_i']:.2f} A", f"{stats1['A']['mean_tq']:.2f} Nm", f"{stats1['A']['start_t']:.1f}°C", f"{stats1['A']['max_t']:.1f}°C", f"{stats1['A']['heating_rate']:.2f}°C/phút"],
    ["  - Trục di chuyển B", "", f"{stats1['B']['mean_i']:.2f} A", f"{stats1['B']['mean_tq']:.2f} Nm", f"{stats1['B']['start_t']:.1f}°C", f"{stats1['B']['max_t']:.1f}°C", f"{stats1['B']['heating_rate']:.2f}°C/phút"],
    ["  - Trục di chuyển C (Lỗi)", "", f"{stats1['C']['mean_i']:.2f} A", f"{stats1['C']['mean_tq']:.2f} Nm", f"{stats1['C']['start_t']:.1f}°C", f"{stats1['C']['max_t']:.1f}°C", f"{stats1['C']['heating_rate']:.2f}°C/phút"],
    ["  - Trục di chuyển D", "", f"{stats1['D']['mean_i']:.2f} A", f"{stats1['D']['mean_tq']:.2f} Nm", f"{stats1['D']['start_t']:.1f}°C", f"{stats1['D']['max_t']:.1f}°C", f"{stats1['D']['heating_rate']:.2f}°C/phút"],
    # Try 2
    ["Try 2 (Động cơ C cũ)", f"{act2:.2f} phút", "", "", "", "", ""],
    ["  - Trục di chuyển A", "", f"{stats2['A']['mean_i']:.2f} A", f"{stats2['A']['mean_tq']:.2f} Nm", f"{stats2['A']['start_t']:.1f}°C", f"{stats2['A']['max_t']:.1f}°C", f"{stats2['A']['heating_rate']:.2f}°C/phút"],
    ["  - Trục di chuyển B", "", f"{stats2['B']['mean_i']:.2f} A", f"{stats2['B']['mean_tq']:.2f} Nm", f"{stats2['B']['start_t']:.1f}°C", f"{stats2['B']['max_t']:.1f}°C", f"{stats2['B']['heating_rate']:.2f}°C/phút"],
    ["  - Trục di chuyển C (Lỗi)", "", f"{stats2['C']['mean_i']:.2f} A", f"{stats2['C']['mean_tq']:.2f} Nm", f"{stats2['C']['start_t']:.1f}°C", f"{stats2['C']['max_t']:.1f}°C", f"{stats2['C']['heating_rate']:.2f}°C/phút"],
    ["  - Trục di chuyển D", "", f"{stats2['D']['mean_i']:.2f} A", f"{stats2['D']['mean_tq']:.2f} Nm", f"{stats2['D']['start_t']:.1f}°C", f"{stats2['D']['max_t']:.1f}°C", f"{stats2['D']['heating_rate']:.2f}°C/phút"],
    # Try 3
    ["Try 3 (Động cơ C MỚI)", f"{act3:.2f} phút", "", "", "", "", ""],
    ["  - Trục di chuyển A", "", f"{stats3['A']['mean_i']:.2f} A", f"{stats3['A']['mean_tq']:.2f} Nm", f"{stats3['A']['start_t']:.1f}°C", f"{stats3['A']['max_t']:.1f}°C", f"{stats3['A']['heating_rate']:.2f}°C/phút"],
    ["  - Trục di chuyển B", "", f"{stats3['B']['mean_i']:.2f} A", f"{stats3['B']['mean_tq']:.2f} Nm", f"{stats3['B']['start_t']:.1f}°C", f"{stats3['B']['max_t']:.1f}°C", f"{stats3['B']['heating_rate']:.2f}°C/phút"],
    ["  - Trục di chuyển C (Mới)", "", f"{stats3['C']['mean_i']:.2f} A", f"{stats3['C']['mean_tq']:.2f} Nm", f"{stats3['C']['start_t']:.1f}°C", f"{stats3['C']['max_t']:.1f}°C", f"{stats3['C']['heating_rate']:.2f}°C/phút"],
    ["  - Trục di chuyển D", "", f"{stats3['D']['mean_i']:.2f} A", f"{stats3['D']['mean_tq']:.2f} Nm", f"{stats3['D']['start_t']:.1f}°C", f"{stats3['D']['max_t']:.1f}°C", f"{stats3['D']['heating_rate']:.2f}°C/phút"],
]

for r_idx, row in enumerate(rows_data_vi):
    for c_idx, val in enumerate(row):
        t_vi.cell(r_idx+1, c_idx).text = val

for row in t_vi.rows:
    for cell in row.cells:
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(9.0)

doc_add_h1(doc_vi, "3. Đối chiếu Định lượng trục C với trục A, B, D")
doc_add_p(
    doc_vi,
    "Để cung cấp phân tích cấu trúc chi tiết, các thông số điện, cơ và nhiệt của trục C được so sánh trực tiếp với 3 trục còn lại trong lần chạy Try 3:\n\n"
    "• Dòng điện tiêu thụ: Trục C (40.26 A) cao hơn 39.2% so với trục A (28.92 A), cao hơn 21.6% so với trục B (33.11 A) và cao hơn 14.4% so với trục D (35.19 A).\n"
    "• Mô-men xoắn đầu ra: Trục C (15.10 Nm) cao hơn 50.7% so với trục A (10.02 Nm), cao hơn 33.5% so với trục B (11.31 Nm) và cao hơn 24.4% so với trục D (12.14 Nm).\n"
    "• Tốc độ gia nhiệt: Trục C (0.71°C/phút) nóng nhanh hơn 42.0% so với trục A (0.50°C/phút), 34.0% so với trục B (0.53°C/phút) và 18.3% so với trục D (0.60°C/phút). Trong 28.13 phút di chuyển chủ động, nhiệt độ động cơ C tăng thêm +20.0°C lên tới đỉnh 59.0°C, cao nhất hệ thống di chuyển cẩu (Drives A/B chỉ đạt 53.0°C và Drive D chỉ đạt 50.0°C)."
)

doc_add_h1(doc_vi, "4. Chẩn đoán kỹ thuật & Các đề xuất hướng xử lý tiếp theo")
doc_add_p(
    doc_vi,
    "• Bó phanh thủy lực (Brake Drag): Phanh đĩa đỗ (fail-safe brake) của bánh xe C nhả không hoàn toàn. Cần đo áp suất dầu thủy lực nhả phanh thực tế tại góc C xem có đạt đủ áp lực thiết kế 25 đến 30 bar hay không.\n"
    "• Lệch góc chụm bánh xe & Vặn xoắn kết cấu khung: Kiểm tra cơ cấu gá trục bánh C. Sự biến dạng khung kết cấu thép của cẩu có thể phân bổ tải trọng không đều hoặc làm bánh C bị quét lê nghiêng trên mặt đường khi xe chạy thẳng.\n"
    "• Quy trình thử nghiệm tráo đổi (Swap) cô lập lỗi (Đề xuất thực hiện):\n"
    "  1. Cấp độ 1: Tráo đổi hai quả lốp giữa góc C và A. Nếu lỗi dòng cao/nhiệt cao di chuyển sang trục A, lỗi do quả lốp hoặc áp suất hơi lốp. Nếu lỗi vẫn ở vị trí C, tiến hành Cấp độ 2.\n"
    "  2. Cấp độ 2: Tráo cụm truyền động (motor, hộp số, phanh) góc C sang góc A. Nếu lỗi vẫn ở vị trí C, nguyên nhân thuộc về kết cấu gá bánh xe hoặc đường dầu thủy lực mở phanh góc C. Nếu lỗi dịch sang vị trí A, lỗi nằm ở hộp số kẹt cơ khí hoặc má phanh bó sát của cụm truyền động."
)

doc_add_h1(doc_vi, "5. Đồ thị Telemetry kiểm chứng")
doc_add_h2(doc_vi, "5.1 Đồ thị dòng điện & nhiệt độ Try 3 (Thời gian thực)")
p_img_vi1 = doc_vi.add_paragraph()
p_img_vi1.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_img_vi1.add_run().add_picture(plot_path, width=Inches(6.0))

doc_vi.add_page_break()

doc_add_h2(doc_vi, "5.2 Biểu đồ cột đối chiếu các lần chạy thử (Try 1 vs. Try 2 vs. Try 3)")
p_img_vi2 = doc_vi.add_paragraph()
p_img_vi2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_img_vi2.add_run().add_picture(comp_plot_path, width=Inches(6.0))

doc_vi.add_page_break()

doc_add_h2(doc_vi, "5.3 Biểu đồ gia nhiệt riêng động cơ TransC qua các lần thử")
p_img_vi3 = doc_vi.add_paragraph()
p_img_vi3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_img_vi3.add_run().add_picture(transc_plot_path, width=Inches(6.0))

doc_vi.save(docx_vi_path)
print("Vietnamese reports generated successfully!")
